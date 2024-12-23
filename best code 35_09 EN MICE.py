import pandas as pd
import numpy as np
import statsmodels.api as sm
from sklearn.linear_model import ElasticNet, ElasticNetCV
from sklearn.preprocessing import StandardScaler
from scipy import stats
from docx import Document

# Load the data from the input Excel file
try:
    df = pd.read_excel('input.xlsx')
except FileNotFoundError:
    print("Error: The file 'input.xlsx' was not found. Please ensure it is in the correct directory.")
    exit()
except Exception as e:
    print(f"Error loading input file: {e}")
    exit()

# Univariate Logistic Regression Analysis
def univariate_logistic_regression(df):
    try:
        results = []
        selected_features = []  # To store features with p-value < 0.2
        for column in df.columns[:-1]:  # All columns except the last one (outcome)
            X = df[column]
            y = df.iloc[:, -1]  # Outcome variable

            # Filter out rows with missing values in X or y
            valid_mask = X.notna() & y.notna()
            X_valid = X[valid_mask]
            y_valid = y[valid_mask]

            # Ensure the outcome is binary
            if not y_valid.isin([0, 1]).all():
                print("Warning: Outcome variable should be binary for logistic regression.")
                return []

            # Perform logistic regression if sufficient data remains
            if len(X_valid) > 1:
                X_valid = sm.add_constant(X_valid)  # Add constant for the intercept
                model = sm.Logit(y_valid, X_valid).fit(disp=False)

                p_value = model.pvalues.iloc[1]  # p-value for the factor
                odds_ratio = np.exp(model.params.iloc[1])  # Odds ratio
                conf = model.conf_int()  # Confidence intervals
                lower_ci = np.exp(conf.iloc[1, 0])  # Lower bound of 95% CI
                upper_ci = np.exp(conf.iloc[1, 1])  # Upper bound of 95% CI

                result = {
                    'Factor': column,
                    'Odds Ratio': odds_ratio,
                    '95% CI Lower': lower_ci,
                    '95% CI Upper': upper_ci,
                    'P-Value': p_value,
                }
                result['Recommendation'] = 'Use this risk factor' if p_value < 0.2 else 'P value >=0.2, do not use'
                results.append(result)

                # Add to selected features if p-value is less than 0.2
                if p_value < 0.2:
                    selected_features.append(column)
            else:
                print(f"Not enough valid data for column: {column}")

        uni_df = pd.DataFrame(results)
        uni_df.to_excel('uni_output.xlsx', index=False)
        print("Univariate Logistic Regression Results:")
        print(uni_df)

        return selected_features  # Return selected features
    except Exception as e:
        print(f"Error during univariate logistic regression: {e}")
        return []


# Elastic Net Regression for Feature Selection
# Elastic Net Regression for Feature Selection
def elastic_net_regression(df, alpha, l1_ratio, selected_features):
    try:
        if not selected_features:
            print("No features selected. Skipping ElasticNet regression.")
            return []

        X = df[selected_features]
        y = df.iloc[:, -1]  # Outcome variable

        # Filter out rows with missing values in X or y
        valid_mask = X.notna().all(axis=1) & y.notna()
        X_valid = X[valid_mask]
        y_valid = y[valid_mask]

        # Standardize the data
        scaler = StandardScaler()
        X_scaled = scaler.fit_transform(X_valid)

        # Perform Elastic Net regression
        elastic_net = ElasticNet(alpha=alpha, l1_ratio=l1_ratio, random_state=42)
        elastic_net.fit(X_scaled, y_valid)

        # Get selected features from ElasticNet
        selected_features_from_en = X.columns[elastic_net.coef_ != 0].tolist()

        # Prepare results for saving
        elastic_net_results = {
            'Feature': X.columns,
            'Coefficient': elastic_net.coef_
        }
        elastic_net_df = pd.DataFrame(elastic_net_results)

        # Save results to an Excel file
        elastic_net_df.to_excel('elastic_output.xlsx', index=False)
        print("ElasticNet regression results saved to 'elastic_output.xlsx'.")

        print("Selected Features after ElasticNet regression:", selected_features_from_en)

        return selected_features_from_en
    except Exception as e:
        print(f"Error during Elastic Net regression: {e}")
        return []


# Function to tune alpha and l1_ratio using ElasticNetCV
def elastic_net_tune_alpha_cv(df):
    try:
        X = df.iloc[:, :-1]
        y = df.iloc[:, -1]

        # Filter out rows with missing values in X or y
        valid_mask = X.notna().all(axis=1) & y.notna()
        X_valid = X[valid_mask]
        y_valid = y[valid_mask]

        # Standardize the data
        scaler = StandardScaler()
        X_scaled = scaler.fit_transform(X_valid)

        # Perform ElasticNetCV over multiple l1_ratios
        l1_ratios = [0.1, 0.5, 0.7, 0.9, 1.0]
        elastic_net_cv = ElasticNetCV(cv=5, l1_ratio=l1_ratios, random_state=42)
        elastic_net_cv.fit(X_scaled, y_valid)

        # Retrieve the best alpha and l1_ratio
        best_alpha = elastic_net_cv.alpha_
        best_l1_ratio = elastic_net_cv.l1_ratio_
        print(f"Best alpha found via ElasticNetCV: {best_alpha}")
        print(f"Best l1_ratio found via ElasticNetCV: {best_l1_ratio}")

        # Save results to a Word document
        doc = Document()
        doc.add_heading('ElasticNetCV Tuning Results', level=1)
        doc.add_paragraph(f"Best Alpha: {best_alpha}")
        doc.add_paragraph(f"Best L1 Ratio: {best_l1_ratio}")
        doc.save('elastic_net_tune.docx')
        print("ElasticNetCV tuning results saved to 'elastic_net_tune.docx'.")

        return best_alpha, best_l1_ratio
    except Exception as e:
        print(f"Error during ElasticNetCV alpha and l1_ratio tuning: {e}")
        return 1.0, 0.5  # Fallback values

# Multivariate Logistic Regression
def multivariate_logistic_regression(df, selected_features):
    try:
        if len(selected_features) == 0:
            print("No features selected. Multivariate logistic regression cannot be performed.")
            return

        X = df[selected_features]
        y = df.iloc[:, -1]

        # Filter out rows with missing values in X or y
        valid_mask = X.notna().all(axis=1) & y.notna()
        X_valid = X[valid_mask]
        y_valid = y[valid_mask]

        X_valid = sm.add_constant(X_valid)
        model = sm.Logit(y_valid, X_valid).fit(disp=False)

        coefficients = model.params
        conf = model.conf_int()
        odds_ratios = np.exp(coefficients)
        lower_ci = np.exp(conf.iloc[:, 0])
        upper_ci = np.exp(conf.iloc[:, 1])

        result_df = pd.DataFrame({
            'Factor': coefficients.index,
            'Odds Ratio': odds_ratios,
            '95% CI Lower': lower_ci,
            '95% CI Upper': upper_ci,
            'P-Value': model.pvalues
        })

        result_df = result_df[result_df['P-Value'] < 0.05]
        if result_df.empty:
            print("No significant features found. Skipping multivariate regression.")
            return

        result_df.to_excel('multi_output.xlsx', index=False)
        print("Multivariate Logistic Regression Results:")
        print(result_df)
    except Exception as e:
        print(f"Error fitting multivariate logistic regression: {e}")

# Hosmer-Lemeshow Test for Goodness of Fit
def hosmer_lemeshow_test(df, selected_features):
    try:
        X = df[selected_features]
        y = df.iloc[:, -1]

        # Filter out rows with missing values
        valid_mask = X.notna().all(axis=1) & y.notna()
        X_valid = X[valid_mask]
        y_valid = y[valid_mask]

        X_valid = sm.add_constant(X_valid)
        model = sm.Logit(y_valid, X_valid).fit(disp=False)

        predicted_probs = model.predict(X_valid)
        observed = y_valid
        expected = predicted_probs

        # Use duplicates='drop' to handle duplicate bin edges in qcut
        deciles = pd.qcut(predicted_probs, 10, labels=False, duplicates='drop')
        observed_freq = observed.groupby(deciles).sum()
        expected_freq = expected.groupby(deciles).sum()

        hl_statistic, p_value = stats.chisquare(f_obs=observed_freq, f_exp=expected_freq)

        hosmer_result = {
            'Hosmer-Lemeshow Statistic': hl_statistic,
            'P-Value': p_value,
            'Conclusion': 'Good fit' if p_value > 0.05 else 'Poor fit'
        }
        hosmer_df = pd.DataFrame([hosmer_result])
        hosmer_df.to_excel('hosmer_test.xlsx', index=False)
        print("Hosmer-Lemeshow Test Results:")
        print(hosmer_df)
    except Exception as e:
        print(f"Error during Hosmer-Lemeshow test: {e}")


# Main program execution
if __name__ == "__main__":
    # Step 1: Univariate Logistic Regression
    selected_univariate_features = univariate_logistic_regression(df)

    # Step 2: ElasticNetCV Tuning
    best_alpha, best_l1_ratio = elastic_net_tune_alpha_cv(df)

    # Step 3: Elastic Net Regression
    selected_features_from_en = elastic_net_regression(df, alpha=best_alpha, l1_ratio=best_l1_ratio, selected_features=selected_univariate_features)

    # Step 4: Multivariate Logistic Regression
    multivariate_logistic_regression(df, selected_features_from_en)

    # Step 5: Hosmer-Lemeshow Test
    hosmer_lemeshow_test(df, selected_features_from_en)
